from docx import Document
import os

try:
    doc = Document()
    doc.add_heading('Bảng I. NỘI DUNG KIỂM TRA / CHECK POINTS (Text Mode)', level=1)
    
    # Simple Paragraph with the text variable
    p = doc.add_paragraph('{{ text_danh_sach_loi }}')
    
    # Add Totals Section
    doc.add_paragraph('--------------------------------------------------')
    doc.add_paragraph('TỔNG CỘNG:')
    doc.add_paragraph('Nặng: {{ tong_loi_nang }}')
    doc.add_paragraph('Nhẹ: {{ tong_loi_nhe }}')
    doc.add_paragraph('Tổng: {{ tong_loi_tong }}')

    out_path = os.path.join("Template", "Template BBK Fix Text.docx")
    doc.save(out_path)
    print(f"Successfully created: {out_path}")

except Exception as e:
    print(f"Error: {e}")
