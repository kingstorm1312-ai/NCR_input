from docx import Document
import os

try:
    doc = Document()
    doc.add_heading('Bảng I. NỘI DUNG KIỂM TRA / CHECK POINTS', level=1)
    
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'STT'
    hdr[1].text = 'Tên lỗi'
    hdr[2].text = 'Số lượng'
    hdr[3].text = 'Vị trí lỗi'
    hdr[4].text = 'Mức độ'
    
    # Content Row
    row = table.rows[1].cells
    
    # Col 1: Loop Start + STT
    row[0].text = '{%tr for i in danh_sach_loi_rut_gon %}{{i.stt}}'
    
    # Col 2: Name
    row[1].text = '{{i.ten_loi}}'
    
    # Col 3: Qty
    row[2].text = '{{i.tong_sl}}'
    
    # Col 4: Detail
    row[3].text = '{{i.chi_tiet}}'
    
    # Col 5: Severity + Loop End
    row[4].text = '{{i.muc_do}}{%tr endfor %}'
    
    # Add Footer Row for Totals
    row_total = table.add_row().cells
    row_total[0].text = 'TỔNG CỘNG'
    # Merge cells if needed, or just put in cols
    row_total[1].text = 'Nặng: {{tong_loi_nang}}'
    row_total[2].text = 'Nhẹ: {{tong_loi_nhe}}'
    row_total[3].text = ''
    row_total[4].text = 'Tổng: {{tong_loi_tong}}'

    out_path = os.path.join("Template", "Template BBK Fix.docx")
    doc.save(out_path)
    print(f"Successfully created: {out_path}")

except Exception as e:
    print(f"Error: {e}")
